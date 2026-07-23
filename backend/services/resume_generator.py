from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)

def add_horizontal_line(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_spacing(paragraph, before=0, after=0):
    pPr = paragraph._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(before))
    pSpacing.set(qn('w:after'), str(after))
    pPr.append(pSpacing)

def add_right_tab(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9720')
    tabs.append(tab)
    pPr.append(tabs)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=180, after=60)
    add_horizontal_line(p)
    run = p.add_run(text.upper())
    set_font(run, size=10.5, bold=True)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=26, after=26)
    run = p.add_run(text)
    set_font(run, size=8.5)
    return p

def add_job_header(doc, title, company, dates):
    p = doc.add_paragraph()
    set_spacing(p, before=80, after=14)
    add_right_tab(p)
    set_font(p.add_run(title), size=9.5, bold=True)
    set_font(p.add_run(" | "), size=9.5, color="555555")
    set_font(p.add_run(company), size=9.5, color="555555")
    set_font(p.add_run("\t"), size=9.5)
    set_font(p.add_run(dates), size=8.5, italic=True, color="888888")

def add_project_header(doc, name, date):
    p = doc.add_paragraph()
    set_spacing(p, before=80, after=14)
    add_right_tab(p)
    set_font(p.add_run(name), size=9.5, bold=True)
    set_font(p.add_run("\t"), size=9.5)
    set_font(p.add_run(date), size=8.5, italic=True, color="888888")

def add_stack_line(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=14)
    set_font(p.add_run(text), size=8, italic=True, color="555555")

def generate_resume(profile, skills, experience, projects, education,
                    output_path, company_name="", role_name="", research:list=None):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=24)
    set_font(p.add_run(profile["name"].upper()), size=17, bold=True)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=50)
    contact = f"{profile['phone']}  |  {profile['email']}  |  {profile['linkedin']}  |  {profile['github']}"
    if profile.get('portfolio'):
        contact += f"  |  {profile['portfolio']}"
    set_font(p.add_run(contact), size=8)

    # Education
    add_section_header(doc, "Education")
    p = doc.add_paragraph()
    set_spacing(p, before=40, after=12)
    set_font(p.add_run(education["ms_degree"]), size=9.5, bold=True)
    set_font(p.add_run(f"  —  {education['ms_school']}"), size=9.5)

    p = doc.add_paragraph()
    set_spacing(p, before=0, after=12)
    set_font(p.add_run(f"Relevant: {education['ms_relevant']}"), size=8, color="555555")

    p = doc.add_paragraph()
    set_spacing(p, before=20, after=40)
    set_font(p.add_run(education["bs_degree"]), size=9.5, bold=True)
    set_font(p.add_run(f"  —  {education['bs_school']}"), size=9.5)

    # Skills — ONE table, forced column widths via XML
    add_section_header(doc, "Technical Skills")

    LABEL_W = 1800  # twips (~1.25 inch)
    VALUE_W = 8280  # twips (~5.75 inch)

    table = doc.add_table(rows=len(skills), cols=2)

    # Set table width
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LABEL_W + VALUE_W))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Remove borders
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Set column widths in tblGrid
    tblGrid = OxmlElement('w:tblGrid')
    for w in [LABEL_W, VALUE_W]:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        tblGrid.append(col)
    tbl.insert(0, tblGrid)

    for i, (label, value) in enumerate(skills.items()):
        row = table.rows[i]

        # Force cell widths
        for cell, width in zip(row.cells, [LABEL_W, VALUE_W]):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # Label
        lp = row.cells[0].paragraphs[0]
        set_spacing(lp, before=30, after=30)
        set_font(lp.add_run(label), size=8.5, bold=True)

        # Value — no wrap by setting font size smaller if needed
        vp = row.cells[1].paragraphs[0]
        set_spacing(vp, before=30, after=30)
        set_font(vp.add_run(value), size=8.5)

    # Experience
    add_section_header(doc, "Professional Experience")
    for job in experience:
        add_job_header(doc, job["title"], job["company"], job["dates"])
        for b in job["bullets"]:
            add_bullet(doc, b)

    # Research (optional — only if research entries exist)
    if research:
        add_section_header(doc, "Research")
        for item in research:
            add_job_header(doc, item["title"], item["institution"], item["dates"])
            for b in item["bullets"]:
                add_bullet(doc, b)

    # Projects
    add_section_header(doc, "Projects")
    for project in projects:
        add_project_header(doc, project["name"], project.get("date", "2026"))
        stack = f"Stack: {', '.join(project['stack'])}"
        if project.get("github_url"):
            stack += f" — {project['github_url']}"
        elif project.get("live_url"):
            stack += f" — {project['live_url']}"
        add_stack_line(doc, stack)
        for b in project["highlights"]:
            add_bullet(doc, b)

    # Accomplishments
    add_section_header(doc, "Accomplishments")
    p = doc.add_paragraph()
    set_spacing(p, before=40, after=14)
    set_font(p.add_run("CometX Accelerator 2026"), size=9.5, bold=True)
    set_font(p.add_run("  —  UT Dallas x Harvard Business School Foundry  |  Top 20 / 181 Teams, Draper Pitch Competition"), size=8.5, color="555555")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


DEFAULT_PROFILE = {
    "name": "Vishwaa Shah",
    "email": "vishwaa.career@gmail.com",
    "phone": "+1 (469) 372-0558",
    "linkedin": "linkedin.com/in/vishwaa-shah",
    "github": "github.com/Vswashah",
    "portfolio": "vishwaashah.vercel.app",
}

DEFAULT_EDUCATION = {
    "ms_degree": "M.S. Computer Science",
    "ms_school": "University of Texas at Dallas  |  Data Science Track",
    "ms_relevant": "Machine Learning, Database Design, Statistical Methods for Data Science, Network Security",
    "bs_degree": "B.Tech, Computer Engineering",
    "bs_school": "Sarvajanik College of Engineering & Technology, India",
}

DEFAULT_EXPERIENCE = [

    {
        "title": "Software Developer Intern",
        "company": "Palm Infotech, Surat, India",
        "dates": "Jan 2025 – Apr 2025",
        "bullets": [
            "Built and maintained production REST APIs in Python/Node.js with MySQL — shipped 10+ features end-to-end in Agile sprints; 85%+ test coverage via GitHub Actions CI/CD with zero regression incidents.",
            "Identified and resolved N+1 database query bottleneck causing production timeouts — added structured logging, rewrote with JOIN and Redis caching; response time dropped from 800ms to <250ms.",
        ]
    },
    {
        "title": "Data Analyst Intern",
        "company": "Technokrit Solutions, India",
        "dates": "May 2024 – Aug 2024",
        "bullets": [
            "Built Power BI dashboards and automated ETL pipelines surfacing KPIs across 3+ business teams — wrote SQL queries to retrieve and transform data from relational databases for business stakeholders.",
            "Applied statistical analysis to identify process optimization opportunities — cleaned and transformed raw datasets, documented data definitions and reporting processes to support consistency across teams.",
        ]
    },
    {
        "title": "Computer Science Grader",
        "company": "University of Texas at Dallas",
        "dates": "Sep 2025 – May 2026",
        "bullets": [
            "Evaluated 300+ student submissions/semester in Data Structures and Algorithms — collaborated with faculty on grading rubrics; maintained detailed feedback logs and academic integrity records.",
            "Co-authored empirical AI security research benchmarking 13 frontier LLMs across 569 real-world GitHub commits using CodeQL static analysis; 426 CPU hours distributed compute; publication-track paper.",
        ]
    },
    {
        "title": "CS Outreach Instructor",
        "company": "University of Texas at Dallas",
        "dates": "Jun 2026 – Present",
        "bullets": [
            "Build and deploy interactive web tools for real student users — take end-to-end ownership of features, iterate based on feedback, deploy via CI/CD pipeline.",
            "Design curriculum and hands-on coding exercises for 20+ students per session — translate complex technical concepts into accessible formats; maintain high engagement across multi-week programs.",
        ]
    },
]

DEFAULT_RESEARCH = [
    {
        "title": "Empirical AI Security Researcher",
        "institution": "University of Texas at Dallas",
        "dates": "Jan 2026 – May 2026",
        "bullets": [
            "Co-authored empirical research benchmarking 13 frontier LLMs across 569 real-world GitHub commits using CodeQL static analysis — designed evaluation framework measuring AI-generated patch quality; 426 CPU hours distributed compute.",
            "Applied statistical methods to analyze LLM performance across security metrics — paper on publication track with PhD reviewer feedback.",
        ]
    },
]


import shutil as _shutil

LIBREOFFICE = (
    _shutil.which("soffice")
    or _shutil.which("libreoffice")
    or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
)

def generate_resume_autofit(profile, skills, experience, projects, education, output_path, company_name="", role_name="", research=None):
    """Auto-fits resume to exactly one page using binary search on spacing multiplier."""
    import fitz
    import os
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)

    def try_mult(m):
        tmp = output_path.replace(".docx", f"_tmp_{m:.3f}.docx")
        tmp_pdf = tmp.replace(".docx", ".pdf")
        try:
            _generate_with_multiplier(profile, skills, experience, projects, education, tmp, research=research, multiplier=m)
            import subprocess
            subprocess.run([LIBREOFFICE, "--headless", "--convert-to", "pdf", "--outdir", output_dir, tmp], capture_output=True)
            doc = fitz.open(tmp_pdf)
            pages = len(doc)
            blocks = doc[0].get_text("blocks")
            remaining = (doc[0].rect.height - max(b[3] for b in blocks)) if pages == 1 and blocks else -1
            doc.close()
            return pages, remaining
        except Exception as e:
            return 2, -1
        finally:
            for f in [tmp, tmp_pdf]:
                try: os.remove(f)
                except: pass

    lo, hi, best = 0.5, 1.0, 0.75
    for _ in range(10):
        mid = (lo + hi) / 2
        pages, rem = try_mult(mid)
        if pages == 1 and rem > 28:  # > 10mm, increase spacing
            best = mid
            lo = mid
        elif pages > 1 or rem < 0:  # overflow, decrease
            hi = mid
        else:  # 1 page, 8-20mm — perfect
            best = mid
            break
    import sys; print(f'Autofit: best multiplier={best:.3f}', file=sys.stderr)

    _generate_with_multiplier(profile, skills, experience, projects, education, output_path, research=research, multiplier=best)
    return output_path


def _generate_with_multiplier(profile, skills, experience, projects, education, output_path, research=None, multiplier=1.0):
    """Generates resume with all spacing scaled by multiplier."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os

    m = multiplier

    def sf(run, size, bold=False, color=None, italic=False):
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
            run.font.color.rgb = RGBColor(r, g, b)

    def sp(p, before=0, after=0):
        pPr = p._p.get_or_add_pPr()
        pSpacing = OxmlElement('w:spacing')
        pSpacing.set(qn('w:before'), str(int(before * m)))
        pSpacing.set(qn('w:after'), str(int(after * m)))
        pPr.append(pSpacing)

    def hline(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def rtab(p):
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9720')
        tabs.append(tab)
        pPr.append(tabs)

    def sec_hdr(text):
        p = doc.add_paragraph()
        sp(p, before=180, after=60)
        hline(p)
        sf(p.add_run(text.upper()), size=10.5, bold=True)

    def blt(text):
        p = doc.add_paragraph(style='List Bullet')
        sp(p, before=26, after=26)
        sf(p.add_run(text), size=8.5)

    def job_hdr(title, company, dates):
        p = doc.add_paragraph()
        sp(p, before=80, after=14)
        rtab(p)
        sf(p.add_run(title), size=9.5, bold=True)
        sf(p.add_run(" | "), size=9.5, color="555555")
        sf(p.add_run(company), size=9.5, color="555555")
        sf(p.add_run("\t"), size=9.5)
        sf(p.add_run(dates), size=8.5, italic=True, color="888888")

    def proj_hdr(name, date):
        p = doc.add_paragraph()
        sp(p, before=80, after=14)
        rtab(p)
        sf(p.add_run(name), size=9.5, bold=True)
        sf(p.add_run("\t"), size=9.5)
        sf(p.add_run(date), size=8.5, italic=True, color="888888")

    def stk(text):
        p = doc.add_paragraph()
        sp(p, before=0, after=14)
        sf(p.add_run(text), size=8, italic=True, color="555555")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.42)
    section.right_margin = Inches(0.42)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=0, after=34)
    sf(p.add_run(profile["name"].upper()), size=17, bold=True)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=0, after=68)
    contact = f"{profile['phone']}  |  {profile['email']}  |  {profile['linkedin']}  |  {profile['github']}"
    if profile.get('portfolio'):
        contact += f"  |  {profile['portfolio']}"
    sf(p.add_run(contact), size=8)

    # Education
    sec_hdr("Education")
    p = doc.add_paragraph()
    sp(p, before=52, after=16)
    sf(p.add_run(education["ms_degree"]), size=9.5, bold=True)
    sf(p.add_run(f"  —  {education['ms_school']}"), size=9.5)
    p = doc.add_paragraph()
    sp(p, before=0, after=16)
    sf(p.add_run(f"Relevant: {education['ms_relevant']}"), size=8, color="555555")
    p = doc.add_paragraph()
    sp(p, before=30, after=52)
    sf(p.add_run(education["bs_degree"]), size=9.5, bold=True)
    sf(p.add_run(f"  —  {education['bs_school']}"), size=9.5)

    # Skills
    sec_hdr("Technical Skills")
    LABEL_W, VALUE_W = 1800, 8100
    table = doc.add_table(rows=len(skills), cols=2)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LABEL_W + VALUE_W))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblGrid = OxmlElement('w:tblGrid')
    for w in [LABEL_W, VALUE_W]:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        tblGrid.append(col)
    tbl.insert(0, tblGrid)
    rs = int(28 * m)
    for i, (label, value) in enumerate(skills.items()):
        row = table.rows[i]
        for cell, width in zip(row.cells, [LABEL_W, VALUE_W]):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        lp = row.cells[0].paragraphs[0]
        sp(lp, before=rs, after=rs)
        sf(lp.add_run(label), size=8.5, bold=True)
        vp = row.cells[1].paragraphs[0]
        sp(vp, before=rs, after=rs)
        sf(vp.add_run(value), size=8.5)

    # Experience
    sec_hdr("Professional Experience")
    for job in experience:
        job_hdr(job["title"], job["company"], job["dates"])
        for b in job["bullets"]:
            blt(b)

    # Research
    if research:
        sec_hdr("Research")
        for item in research:
            job_hdr(item["title"], item["institution"], item["dates"])
            for b in item["bullets"]:
                blt(b)

    # Projects
    sec_hdr("Projects")
    for project in projects:
        proj_hdr(project["name"], project.get("date", "2026"))
        s = f"Stack: {', '.join(project['stack'])}"
        if project.get("github_url"):
            s += f" — {project['github_url']}"
        elif project.get("live_url"):
            s += f" — {project['live_url']}"
        stk(s)
        for b in project["highlights"]:
            blt(b)

    # Accomplishments
    sec_hdr("Accomplishments")
    p = doc.add_paragraph()
    sp(p, before=52, after=12)
    sf(p.add_run("CometX Accelerator 2026"), size=9.5, bold=True)
    sf(p.add_run("  —  UT Dallas x Harvard Business School Foundry  |  Top 20 / 181 Teams, Draper Pitch Competition"), size=8.5, color="555555")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
